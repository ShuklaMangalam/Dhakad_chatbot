import streamlit as st
import os
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv


load_dotenv()
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")
os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGCHAIN_API_KEY")
os.environ["LANGCHAIN_TRACING_V2"] = "True"


llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")


prompt_template = ChatPromptTemplate.from_messages([
    ("system", "You are a helpful assistant. Please provide responses to user queries."),
    ("user", "Question: {question}")
])


output_parser = StrOutputParser()


st.set_page_config(page_title="🧠 Dhakad AI Assistant", layout="centered")
st.title("🧠 Dhakad Chatbot with Document Reader")
st.markdown("Ask general questions or upload a file and ask about its content!")


input_text = st.text_input("💬 Ask a general question:")
if input_text:
    chain = prompt_template | llm | output_parser
    response = chain.invoke({"question": input_text})
    st.markdown("### 🤖 Gemini's Answer:")
    st.write(response)
    st.snow()

st.markdown("---")


uploaded_file = st.file_uploader("📁 Upload a file", type=["pdf", "docx", "csv", "xlsx"])


if "db" not in st.session_state:
    st.session_state.db = None
    st.session_state.file_type = None


if uploaded_file:
    file_name = uploaded_file.name
    file_type = file_name.split(".")[-1].lower()

    st.session_state.file_type = file_type

    with st.spinner(f"📄 Processing {file_type.upper()} file..."):

        
        temp_file_path = f"temp_{uploaded_file.name}"
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getvalue())

        
        if file_type == "pdf":
            loader = PyPDFLoader(temp_file_path)
            docs = loader.load()

        elif file_type == "docx":
            doc = Document(temp_file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            from langchain.docstore.document import Document as LangchainDocument
            docs = [LangchainDocument(page_content=text)]

        elif file_type == "csv":
            df = pd.read_csv(temp_file_path)
            text = df.to_string()
            from langchain.docstore.document import Document as LangchainDocument
            docs = [LangchainDocument(page_content=text)]

        elif file_type == "xlsx":
            df = pd.read_excel(temp_file_path)
            text = df.to_string()
            from langchain.docstore.document import Document as LangchainDocument
            docs = [LangchainDocument(page_content=text)]

        else:
            st.error("❌ Unsupported file type.")
            os.remove(temp_file_path)
            st.stop()

        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        texts = text_splitter.split_documents(docs)


        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        db = FAISS.from_documents(texts, embeddings)

        
        st.session_state.db = db

        
        os.remove(temp_file_path)

        st.success(f"✅ {file_type.upper()} file processed successfully!")


input_query = st.text_input("🔍 Ask a question about the uploaded file:")


if input_query:
    if st.session_state.db is not None:
        with st.spinner("🧠 DHAKAD is thinking..."):
            retriever = st.session_state.db.as_retriever(search_kwargs={"k": 3})
            docs = retriever.invoke(input_query)

            
            context = "\n".join([doc.page_content for doc in docs])

            
            full_prompt = f"""
            Context: {context}

            Question: {input_query}
            Answer:
            """

            
            chain = prompt_template | llm | output_parser
            response = chain.invoke({"question": full_prompt})

            st.markdown("### 🧠 DHAKAD's Answer:")
            st.write(response)
            st.snow()
    else:
        st.warning("⚠️ Please upload a file first before asking questions.")
